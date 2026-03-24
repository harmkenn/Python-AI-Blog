import streamlit as st
from docx import Document
from docx.oxml.shared import OxmlElement, qn
from io import BytesIO
import mammoth

st.title("📄 Auto‑Format Word Document (Left‑Justify Images + Tight Wrap)")

uploaded_file = st.file_uploader("Upload your .docx file", type=["docx"])

def left_wrap_image(run):
    """Apply left justification + square wrap to an image."""
    inline = run._r.xpath("w:drawing/wp:inline")
    if not inline:
        return

    inline = inline[0]

    # Convert inline to anchor (required for text wrapping)
    inline.tag = qn("wp:anchor")

    # Add wrapSquare
    wrap = OxmlElement("wp:wrapSquare")
    wrap.set("wrapText", "bothSides")
    inline.append(wrap)

    # Position image on the left
    position_h = OxmlElement("wp:positionH")
    position_h.set("relativeFrom", "column")
    align = OxmlElement("wp:align")
    align.text = "left"
    position_h.append(align)
    inline.append(position_h)

def tighten_paragraph_spacing(paragraph):
    """Reduce spacing before/after paragraphs."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()

    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "60")  # 6 pt
    pPr.append(spacing)

if uploaded_file:
    doc = Document(uploaded_file)
    new_doc = Document()

    for para in doc.paragraphs:
        new_para = new_doc.add_paragraph()
        tighten_paragraph_spacing(new_para)

        for run in para.runs:
            new_run = new_para.add_run(run.text)

            # If run contains an image, reformat it
            if "graphic" in run._r.xml:
                left_wrap_image(new_run)

    # Save to buffer
    buffer = BytesIO()
    new_doc.save(buffer)
    buffer.seek(0)

    st.success("Formatting complete! Scroll down to preview your document.")

    # Convert to HTML for preview
    html = mammoth.convert_to_html(buffer).value

    st.markdown("### 📘 Document Preview")
    st.components.v1.html(html, height=600, scrolling=True)

    # Download button
    st.download_button(
        label="⬇️ Download formatted .docx",
        data=buffer,
        file_name="formatted_document.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
